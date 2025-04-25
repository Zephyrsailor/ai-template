"""
Word文档结构化解析模块.
直接从Word文档中提取标题和内容结构，实现基于文档层级的切割。
"""
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple

logger = logging.getLogger(__name__)

def parse_docx_structure(file_path: str) -> List[Dict[str, Any]]:
    """
    解析Word文档的结构，按照标题层级切割内容，生成三部分结构：
    1. 文档标题
    2. 目录
    3. 各章节（标题+内容）
    
    Args:
        file_path: Word文档路径
        
    Returns:
        结构化的文档块列表
    """
    try:
        import docx
        doc = docx.Document(file_path)
        file_name = Path(file_path).name
        blocks = []
        
        # 识别文档标题(通常是第一段)
        main_title = ""
        if doc.paragraphs and doc.paragraphs[0].text.strip():
            main_title = doc.paragraphs[0].text.strip()
            # 确认这是标题而不是其他内容
            if len(main_title) < 100 and not main_title.startswith("注："):
                blocks.append({
                    "type": "document_title",
                    "level": 0,
                    "text": main_title,
                    "file_name": file_name
                })
        
        # 查找并提取目录 - 直接读取整个目录页
        toc_text, toc_paragraphs = extract_toc(doc)
        has_real_toc = False
        
        if toc_text:
            has_real_toc = True
            blocks.append({
                "type": "toc",
                "level": 0,
                "text": toc_text,
                "file_name": file_name
            })
            
            # 构建目录映射，用于后续章节标题
            toc_items = extract_toc_items(toc_paragraphs)
            logger.info(f"提取到 {len(toc_items)} 个目录项")
        else:
            toc_items = []
            logger.info("未检测到目录部分")
        
        # 解析正文章节 - 修改为一级标题+内容的方式
        chapters = extract_chapters(doc, file_name, toc_items)
        blocks.extend(chapters)
        
        # 如果没有识别出任何块，创建一个默认块
        if not blocks:
            logger.warning(f"未能从文档中提取结构: {file_path}")
            # 获取文档全文
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            blocks.append({
                "type": "content",
                "level": 0,
                "text": full_text,
                "file_name": file_name
            })
        
        # 如果没有目录但有章节，创建一个虚拟目录
        if not has_real_toc and chapters:
            virtual_toc_text = "目录\n\n"
            
            # 按照原文档格式构建虚拟目录
            for chapter in chapters:
                if chapter["type"] == "chapter":
                    chapter_title = chapter.get("title", "")
                    chapter_level = chapter.get("level", 1)
                    
                    if chapter_title:
                        # 根据级别添加缩进和格式
                        if chapter_level == 1:
                            if re.match(r'^[一二三四五六七八九十]+、', chapter_title):
                                virtual_toc_text += f"{chapter_title}\n"
                            else:
                                virtual_toc_text += f"{chapter_title}\n" 
                        elif chapter_level == 2:
                            virtual_toc_text += f"    {chapter_title}\n"
                        elif chapter_level == 3:
                            virtual_toc_text += f"        {chapter_title}\n"
                        else:
                            virtual_toc_text += f"{chapter_title}\n"
            
            # 只有当虚拟目录有实际内容时才添加
            if len(virtual_toc_text.split("\n")) > 3:
                blocks.insert(1 if main_title else 0, {
                    "type": "toc",
                    "level": 0,
                    "text": virtual_toc_text,
                    "is_virtual": True,
                    "file_name": file_name
                })
                logger.info("已生成虚拟目录")
        
        return blocks
    except ImportError:
        logger.error("未安装python-docx库，无法解析Word文档结构")
        return []
    except Exception as e:
        logger.exception(f"解析Word文档时出错: {e}")
        return []

def extract_toc_items(toc_paragraphs: list) -> Dict[str, Dict[str, Any]]:
    """
    从目录段落中提取目录项信息
    
    Args:
        toc_paragraphs: 目录部分的段落列表
        
    Returns:
        目录项映射字典：{标题文本: {"level": 级别, "number": 编号}}
    """
    toc_items = {}
    
    for para in toc_paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # 移除页码部分以获取纯标题
        title_text = re.sub(r'\s+\d+$', '', text).strip()
        
        # 检测标题级别
        level = 1  # 默认为一级标题
        
        # 根据开头的编号模式判断级别 
        if re.match(r'^[一二三四五六七八九十]+、', title_text):
            level = 1
            number = re.match(r'^([一二三四五六七八九十]+)、', title_text).group(1)
        elif re.match(r'^\d+\.\s+', title_text):
            level = 2
            number = re.match(r'^(\d+)\.', title_text).group(1)
        elif re.match(r'^\(\d+\)\s+', title_text) or re.match(r'^（\d+）\s+', title_text):
            level = 3
            number = re.search(r'\((\d+)\)|\（(\d+)\）', title_text).group(1) if re.search(r'\((\d+)\)', title_text) else re.search(r'\（(\d+)\）', title_text).group(1)
        else:
            # 无法确定级别的情况
            level = 1
            number = ""
        
        toc_items[title_text] = {
            "level": level,
            "number": number
        }
        
    return toc_items

def extract_toc(doc) -> Tuple[str, List]:
    """
    提取或生成文档的目录部分，确保生成结构清晰的目录文本
    
    Args:
        doc: docx文档对象
        
    Returns:
        (目录文本内容, 目录段落对象列表)
    """
    # 首先尝试找到原始目录
    start_idx = -1
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        if text == "目录" or text == "目 录" or text.lower() == "contents" or text.lower() == "table of contents":
            start_idx = i
            break
    
    # 如果找到了目录标记，尝试提取原始目录
    if start_idx >= 0:
        try:
            # 提取原始目录内容
            toc_text = "目录\n\n"
            toc_paragraphs = [doc.paragraphs[start_idx]]
            end_idx = min(start_idx + 50, len(doc.paragraphs))
            
            for i in range(start_idx + 1, end_idx):
                text = doc.paragraphs[i].text.strip()
                if not text:
                    continue
                    
                # 检查是否已经到了目录结尾（例如遇到了章节开始）
                if (i > start_idx + 5 and  # 确保至少处理了几个目录项
                    (re.match(r'^第[一二三四五六七八九十\d]+[章节篇]', text) or
                     re.match(r'^[一二三四五六七八九十]+、', text)) and
                    not text.endswith(tuple('0123456789'))):  # 确保不是目录项
                    break
                
                # 处理目录项并添加到结果中
                toc_text += format_toc_line(text) + "\n"
                toc_paragraphs.append(doc.paragraphs[i])
                
            # 如果找到了足够的目录项，返回原始目录
            if len(toc_paragraphs) > 3:
                return toc_text.strip(), toc_paragraphs
        except Exception as e:
            logger.warning(f"提取原始目录时出错: {e}，将生成新目录")
    
    # 如果没有找到原始目录或提取失败，通过分析文档结构生成新目录
    return generate_toc_from_structure(doc)

def format_toc_line(text):
    """
    格式化目录行，保持原始的缩进和结构，美化页码对齐
    
    Args:
        text: 原始目录行文本
        
    Returns:
        格式化后的目录行
    """
    # 确定缩进级别
    indent = ""
    
    # 一级标题模式
    if re.match(r'^[一二三四五六七八九十]+、', text) or re.match(r'^第[一二三四五六七八九十\d]+[章节篇]', text):
        indent = ""
    # 二级标题模式
    elif re.match(r'^\d+\.\s+', text) or re.match(r'^\d+\s+', text):
        indent = "    "
    # 三级标题模式
    elif re.match(r'^\(\d+\)\s+', text) or re.match(r'^（\d+）\s+', text):
        indent = "        "
    # 四级标题或其他
    else:
        # 根据前导空格推断缩进
        leading_spaces = len(text) - len(text.lstrip())
        if leading_spaces > 12:
            indent = "            "  # 四级缩进
        elif leading_spaces > 8:
            indent = "        "      # 三级缩进
        elif leading_spaces > 4:
            indent = "    "          # 二级缩进
    
    # 处理页码对齐
    if re.search(r'\d+$', text):
        # 提取页码部分
        match = re.search(r'(\d+)$', text)
        if match:
            page_num = match.group(1)
            title_part = text[:match.start()].strip()
            
            # 如果标题和页码之间有省略号或点，保留它们
            if "..." in title_part:
                parts = title_part.split("...")
                title_part = parts[0].strip()
                
            # 创建点状填充并右对齐页码
            dots = "." * max(3, 60 - len(indent) - len(title_part) - len(page_num) - 2)
            return f"{indent}{title_part} {dots} {page_num}"
    
    # 如果没有页码或不需要特殊处理，保持原样并添加缩进
    # 移除原有的前导空格以避免重复缩进
    clean_text = text.lstrip()
    return f"{indent}{clean_text}"

def generate_toc_from_structure(doc) -> Tuple[str, List]:
    """
    通过分析文档结构生成目录
    
    Args:
        doc: docx文档对象
        
    Returns:
        (目录文本内容, 目录段落对象列表)
    """
    structure = []
    doc_title = ""
    
    # 跳过文档标题(通常是第一段)
    start_idx = 0
    if doc.paragraphs and doc.paragraphs[0].text.strip():
        doc_title = doc.paragraphs[0].text.strip()
        if len(doc_title) < 100 and not doc_title.startswith("注："):
            start_idx = 1
    
    # 分析文档结构，识别标题及其级别
    for i in range(start_idx, min(len(doc.paragraphs), 500)):  # 限制处理范围
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        if not text:
            continue
            
        # 识别是否是标题
        level, is_heading = identify_heading(para, text)
        
        if is_heading:
            structure.append({
                "text": text,
                "level": level,
                "index": i,
                "paragraph": para
            })
    
    # 生成目录文本
    toc_text = "目录\n\n"
    page_number = 1
    
    for item in structure:
        indent = "    " * item["level"]
        text = item["text"]
        
        # 为每个标题分配页码
        if item["level"] == 0:  # 一级标题增加页码
            page_number += 1
        page = str(page_number)
        
        # 格式化目录项，添加点线和页码
        dots = "." * max(3, 60 - len(indent) - len(text) - len(page) - 2)
        toc_text += f"{indent}{text} {dots} {page}\n"
    
    # 创建一个虚拟的段落列表，用于返回
    toc_paragraphs = [doc.paragraphs[0]] if doc.paragraphs else []
    
    return toc_text.strip(), toc_paragraphs

def identify_heading(para, text: str, toc_items: Dict[str, Dict[str, Any]] = None) -> Tuple[int, bool]:
    """
    识别段落是否为标题及其级别
    
    Args:
        para: 段落对象
        text: 段落文本
        toc_items: 目录项映射字典（可选）
        
    Returns:
        (标题级别, 是否为标题)
        标题级别：0=一级标题, 1=二级标题, 2=三级标题
    """
    # 首先检查是否在目录项中
    if toc_items and text in toc_items:
        return (toc_items[text]["level"], True)
    
    # 尝试匹配不带页码的目录项
    if toc_items:
        for toc_title, info in toc_items.items():
            clean_title = re.sub(r'\s+\d+$', '', toc_title).strip()
            if text == clean_title or text.startswith(clean_title):
                return (info["level"], True)
    
    # 检查段落样式
    style_level = 0
    if hasattr(para, 'style') and para.style and para.style.name.startswith('Heading'):
        try:
            style_level = int(para.style.name.split()[-1])
            return (style_level - 1, True)  # 转为0-based
        except:
            pass
    
    # 检查字体特征(粗体/大字体通常是标题)
    is_emphasized = False
    for run in para.runs:
        if run.bold or (hasattr(run, 'font') and hasattr(run.font, 'size') and 
                        run.font.size and run.font.size > 12):
            is_emphasized = True
            break
    
    # 基于文本模式识别标题级别
    if re.match(r"^[一二三四五六七八九十]+、\s*", text):
        return (0, True)  # 一级标题: 一、二、三、
    elif re.match(r"^第[一二三四五六七八九十\d]+[章节篇]", text):
        return (0, True)  # 一级标题: 第一章、第二章
    elif re.match(r"^\d+\.\s+", text) and len(text) < 100:
        return (1, True)  # 二级标题: 1. 2. 3.
    elif (re.match(r"^\(\d+\)\s+", text) or re.match(r"^（\d+）\s+", text)) and len(text) < 100:
        return (2, True)  # 三级标题: (1) (2) (3) or （1） （2）
    
    # 额外检查: 短文本且强调格式可能是标题
    if is_emphasized and len(text) < 50 and not text.endswith("。"):
        return (0, True)
    
    # 进一步基于内容特征检查
    # 标题通常较短，不包含终止符号如句号
    if len(text) < 30 and not text.endswith("。") and any(char.isdigit() for char in text):
        return (1, True)
    
    return (0, False)  # 非标题

def extract_chapters(doc, file_name: str, toc_items: Dict[str, Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """
    提取文档的章节内容，每个一级标题及其内容作为一个章节
    
    Args:
        doc: docx文档对象
        file_name: 文件名
        toc_items: 目录项映射字典
        
    Returns:
        章节块列表
    """
    chapters = []
    
    # 识别所有段落的标题级别
    para_infos = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            para_infos.append({"text": "", "is_heading": False, "level": 0})
            continue
            
        heading_level, is_heading = identify_heading(para, text, toc_items)
        para_infos.append({
            "text": text,
            "is_heading": is_heading,
            "level": heading_level
        })
    
    # 组织成章节结构
    chapter_title = ""
    chapter_content = ""
    chapter_level = 0
    in_chapter = False
    previous_headings = []  # 用于跟踪前一个标题的信息
    
    for i, para_info in enumerate(para_infos):
        text = para_info["text"]
        is_heading = para_info["is_heading"]
        level = para_info["level"]
        
        # 跳过空文本
        if not text:
            if in_chapter and chapter_content:
                chapter_content += "\n"
            continue
        
        # 找到一级标题，开始新章节
        if is_heading and level == 1:
            # 保存先前的章节
            if in_chapter and chapter_title:
                # 只有当章节有实际内容时才保存
                if chapter_content.strip() and len(chapter_content.strip()) > len(chapter_title):
                    chapters.append({
                        "type": "chapter",
                        "level": chapter_level,
                        "title": chapter_title,
                        "text": chapter_title + "\n\n" + chapter_content.strip(),
                        "file_name": file_name
                    })
                elif previous_headings:
                    # 如果前一个章节只有标题没有内容，附加到其他章节
                    previous_chapter = previous_headings[-1] if previous_headings else None
                    if previous_chapter:
                        chapter_title = previous_chapter["title"]
                        chapter_content = previous_chapter["content"]
                        chapter_level = previous_chapter["level"]
            
            # 记录当前章节标题信息
            previous_headings.append({
                "title": text,
                "level": level,
                "content": "",
                "index": i
            })
            
            # 开始新章节
            chapter_title = text
            chapter_content = ""
            chapter_level = level
            in_chapter = True
            
        # 二级或三级标题
        elif is_heading and level > 1 and in_chapter:
            # 在内容中添加标题，而不是开始新章节
            chapter_content += "\n\n" + text + "\n"
            
            # 记录子标题信息
            previous_headings.append({
                "title": text,
                "level": level,
                "content": "",
                "parent_title": chapter_title,
                "index": i
            })
            
        # 在章节内的普通内容
        elif in_chapter:
            chapter_content += text + "\n"
            
            # 更新最近标题的内容
            if previous_headings:
                previous_headings[-1]["content"] += text + "\n"
    
    # 保存最后一个章节
    if in_chapter and chapter_title:
        # 只有当章节有实际内容时才保存
        if chapter_content.strip() and len(chapter_content.strip()) > len(chapter_title):
            chapters.append({
                "type": "chapter",
                "level": chapter_level,
                "title": chapter_title,
                "text": chapter_title + "\n\n" + chapter_content.strip(),
                "file_name": file_name
            })
    
    # 如果章节很少但有很多内容，尝试合并孤立内容
    if len(chapters) < 3:
        merged_content = ""
        for info in para_infos:
            if info["text"]:
                merged_content += info["text"] + "\n\n"
        
        if merged_content and len(merged_content) > 200:
            title = "文档内容"
            if chapters and chapters[0]["title"]:
                title = chapters[0]["title"]
            
            # 替换所有章节为一个合并章节
            chapters = [{
                "type": "chapter",
                "level": 1,
                "title": title,
                "text": merged_content.strip(),
                "file_name": file_name
            }]
    
    # 如果没有发现标准章节，尝试处理整个文档内容
    if not chapters:
        # 跳过标题和目录部分
        start_idx = 0
        for i, para_info in enumerate(para_infos):
            if "目录" in para_info["text"] and i < 30:
                # 找到目录后的第一个非空内容作为起点
                for j in range(i+1, len(para_infos)):
                    if para_infos[j]["text"] and len(para_infos[j]["text"]) > 50:
                        start_idx = j
                        break
                break
        
        # 收集所有剩余内容
        all_content = "\n".join([p["text"] for p in para_infos[start_idx:] if p["text"]])
        
        if all_content:
            chapters.append({
                "type": "chapter",
                "level": 1,
                "title": "文档内容",
                "text": all_content,
                "file_name": file_name
            })
    
    # 确保所有章节都有实际内容
    valid_chapters = []
    for chapter in chapters:
        content_without_title = chapter["text"].replace(chapter["title"], "").strip()
        if len(content_without_title) > 50:  # 至少50个字符的实际内容
            valid_chapters.append(chapter)
        elif valid_chapters:
            # 如果章节内容太少，合并到前一个章节
            prev_chapter = valid_chapters[-1]
            prev_chapter["text"] += "\n\n" + chapter["text"]
    
    # 如果有效章节为空，使用整个文档内容
    if not valid_chapters and para_infos:
        all_text = "\n".join([p["text"] for p in para_infos if p["text"]])
        if all_text:
            valid_chapters.append({
                "type": "chapter",
                "level": 1,
                "title": "文档内容",
                "text": all_text,
                "file_name": file_name
            })
    
    return valid_chapters 