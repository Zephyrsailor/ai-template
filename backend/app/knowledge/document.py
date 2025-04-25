"""
Document loading and handling.
"""
import logging
import re
from pathlib import Path
from typing import Optional, List, Dict, Any
import subprocess
import shutil
from llama_index.core.node_parser import SentenceSplitter
from .docx_parser import parse_docx_structure
from .chunking import StructureAwareChunker

# 使用基本导入方式，避免错误
from llama_index.core.schema import Document as LlamaDocument, TextNode
from llama_index.core.readers import SimpleDirectoryReader
from unstructured.partition.auto import partition

# 尝试导入UnstructuredReader
try:
    from llama_index.readers.file import UnstructuredReader
    UNSTRUCTURED_AVAILABLE = True
    # 不使用可能不支持的参数
    unstructured_reader = UnstructuredReader()
    logger = logging.getLogger(__name__)
    logger.info("成功加载UnstructuredReader")
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    unstructured_reader = None
    logger = logging.getLogger(__name__)
    logger.warning("无法加载UnstructuredReader，将使用基本文本提取方法")

# 尝试导入 mistune 用于 Markdown 解析
try:
    import mistune
except ImportError:
    mistune = None

logger = logging.getLogger(__name__)

import os
import csv
import json
from dataclasses import dataclass

@dataclass
class Document:
    """表示知识库中的文档"""
    text: str
    metadata: Optional[Dict[str, Any]] = None
    score: float = 0.0

def load_documents_from_file(file_path: str) -> List[Document]:
    """
    从文件加载文档
    
    Args:
        file_path: 文件路径
        
    Returns:
        Document列表
    """
    file_path_obj = Path(os.path.abspath(file_path))
    file_extension = file_path_obj.suffix.lower()
    
    try:
        # 根据文件类型选择合适的加载方法
        if file_extension == '.csv':
            return load_from_csv(str(file_path_obj))
        elif file_extension == '.json':
            return load_from_json(str(file_path_obj))
        elif file_extension == '.txt':
            # 修改：调用 load_document 以处理纯文本
            return load_document(file_path_obj, parsing_method='direct') 
        elif file_extension in ['.md', '.markdown']:
            # 修改：调用 load_document 以处理Markdown
            return load_document(file_path_obj, parsing_method='direct') 
        elif file_extension == '.docx':
            # 修改：明确调用 load_document 并强制使用 Pandoc 方法
            logger.info(f"检测到DOCX文件，将使用 Pandoc 方法加载: {file_path_obj}")
            return load_document(file_path_obj, parsing_method='pandoc')
        elif file_extension == '.doc' or file_extension == '.pdf': 
            # 对于 .doc 和 .pdf，仍然尝试使用 unstructured (如果可用)
            if UNSTRUCTURED_AVAILABLE:
                return load_from_unstructured(str(file_path_obj))
            else:
                logger.warning(f"无法处理文件 {file_path_obj}，未安装UnstructuredReader")
                return []
        else:
            logger.warning(f"不支持的文件格式: {file_extension}")
            return []
    except Exception as e:
        logger.error(f"加载文件时出错 {file_path_obj}: {str(e)}")
        return []

def load_from_csv(file_path: str) -> List[Document]:
    """从CSV文件加载文档"""
    documents = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            headers = next(reader, None)
            
            if not headers:
                logger.warning(f"CSV文件为空或缺少标题行: {file_path}")
                return []
            
            # 确定文本列和元数据列
            text_col = 0  # 默认第一列是文本
            metadata_cols = {}
            for i, header in enumerate(headers):
                if i != text_col:
                    metadata_cols[header] = i
            
            for row in reader:
                if not row:
                    continue
                
                # 确保行有足够的列
                if len(row) <= text_col:
                    logger.warning(f"跳过格式不正确的行: {row}")
                    continue
                
                text = row[text_col]
                
                # 构建元数据
                metadata = {}
                for header, col_idx in metadata_cols.items():
                    if col_idx < len(row) and row[col_idx]:
                        metadata[header] = row[col_idx]
                
                # 创建文档
                doc = Document(text=text, metadata=metadata)
                documents.append(doc)
        
        logger.info(f"从CSV文件加载了 {len(documents)} 个文档")
        return documents
    except Exception as e:
        logger.error(f"加载CSV文件时出错 {file_path}: {str(e)}")
        return []

def load_from_json(file_path: str) -> List[Document]:
    """从JSON文件加载文档"""
    documents = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
            # 处理不同的JSON格式
            if isinstance(data, list):
                # 列表格式，每个元素是一个文档
                for item in data:
                    if isinstance(item, dict):
                        # 确定文本字段
                        text = None
                        for key in ['text', 'content', 'body', 'description']:
                            if key in item:
                                text = item.pop(key)
                                break
                        
                        if text:
                            # 其余字段作为元数据
                            doc = Document(text=text, metadata=item)
                            documents.append(doc)
            
            elif isinstance(data, dict):
                # 字典格式，可能包含多个文档
                if 'documents' in data and isinstance(data['documents'], list):
                    # 标准格式：{'documents': [{...}, {...}]}
                    for item in data['documents']:
                        if isinstance(item, dict) and 'text' in item:
                            text = item.pop('text')
                            doc = Document(text=text, metadata=item)
                            documents.append(doc)
                else:
                    # 单个文档格式
                    text = None
                    for key in ['text', 'content', 'body', 'description']:
                        if key in data:
                            text = data.pop(key)
                            break
                    
                    if text:
                        doc = Document(text=text, metadata=data)
                        documents.append(doc)
        
        logger.info(f"从JSON文件加载了 {len(documents)} 个文档")
        return documents
    except Exception as e:
        logger.error(f"加载JSON文件时出错 {file_path}: {str(e)}")
        return []

def load_from_text(file_path: str) -> List[Document]:
    """从文本文件加载文档"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # 创建基本元数据
        metadata = {
            'source': os.path.basename(file_path),
            'created_at': '2025-04-07',
            'content_type': '内容'
        }
        
        # 将文本分块
        documents = []
        # 简单的段落分块
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        
        for i, paragraph in enumerate(paragraphs):
            if len(paragraph) > 20:  # 跳过非常短的段落
                doc = Document(
                    text=paragraph,
                    metadata=metadata.copy()
                )
                documents.append(doc)
        
        logger.info(f"从文本文件加载了 {len(documents)} 个文档")
        return documents
    except Exception as e:
        logger.error(f"加载文本文件时出错 {file_path}: {str(e)}")
        return []

def load_from_markdown(file_path: str) -> List[Document]:
    """从Markdown文件加载文档"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # 创建基本元数据
        metadata = {
            'source': os.path.basename(file_path),
            'created_at': '2025-04-07',
            'content_type': '内容'
        }
        
        # 将文本分块：按标题分块
        documents = []
        lines = text.split('\n')
        current_title = ""
        current_content = []
        
        for line in lines:
            if line.startswith('#'):  # 检测标题行
                # 保存之前的块
                if current_content:
                    content = '\n'.join(current_content)
                    if len(content) > 20:  # 跳过非常短的内容
                        doc_metadata = metadata.copy()
                        if current_title:
                            doc_metadata['title'] = current_title
                        
                        doc = Document(
                            text=content,
                            metadata=doc_metadata
                        )
                        documents.append(doc)
                
                # 开始新块
                current_title = line.lstrip('#').strip()
                current_content = []
            else:
                current_content.append(line)
        
        # 保存最后一个块
        if current_content:
            content = '\n'.join(current_content)
            if len(content) > 20:
                doc_metadata = metadata.copy()
                if current_title:
                    doc_metadata['title'] = current_title
                
                doc = Document(
                    text=content,
                    metadata=doc_metadata
                )
                documents.append(doc)
        
        logger.info(f"从Markdown文件加载了 {len(documents)} 个文档")
        return documents
    except Exception as e:
        logger.error(f"加载Markdown文件时出错 {file_path}: {str(e)}")
        return []

def load_from_unstructured(file_path: str) -> List[Document]:
    """使用unstructured库加载结构化文档"""
    if not partition:
        logger.error("未安装unstructured库，无法加载结构化文档")
        return []
    
    try:
        logger.info("使用UnstructuredReader加载文档")
        elements = partition(file_path)
        
        # --- 添加调试日志 ---
        logger.debug(f"Unstructured partition elements ({len(elements)}):")
        for i, element in enumerate(elements):
             # 打印元素的类别和前100个字符的文本
             element_text_preview = str(element)[:100].replace("\n", " ")
             logger.debug(f"  Element {i}: Category='{element.category}', Text='{element_text_preview}...'" )
        # --- 结束调试日志 ---
        
        # 创建基本元数据
        metadata = {
            'source': os.path.basename(file_path),
            'created_at': '2025-04-07',
            'content_type': '内容'
        }
        
        documents = []
        for element in elements:
            # 提取元素文本和元数据
            element_text = str(element)
            if len(element_text) < 20:  # 跳过非常短的元素
                logger.debug(f"  Skipping short element: {element_text_preview}...")
                continue
                
            # 创建文档
            element_metadata = metadata.copy()
            element_metadata['element_type'] = element.category # 使用unstructured的类别
            
            # --- 添加元数据创建日志 ---
            logger.debug(f"    Creating Document with metadata: {element_metadata}")
            # --- 结束元数据创建日志 ---
            
            doc = Document(
                text=element_text,
                metadata=element_metadata
            )
            documents.append(doc)
        
        logger.info(f"成功加载UnstructuredReader")
        logger.info(f"从结构化文档加载了 {len(documents)} 个文档")
        return documents
    except Exception as e:
        logger.error(f"使用unstructured加载文档时出错 {file_path}: {str(e)}")
        return []

def load_document(file_path: Path, parsing_method: str = 'direct') -> List[Document]:
    """
    加载文档并根据选择的方法进行解析
    
    Args:
        file_path: 文档文件的路径
        parsing_method: 解析方法 ('direct' 或 'pandoc')
        
    Returns:
        Document对象列表
        
    Raises:
        ValueError: 如果文件类型不支持
        FileNotFoundError: 如果文件不存在
    """
    if not file_path.exists():
        raise FileNotFoundError(f"文件未找到: {file_path}")

    file_ext = file_path.suffix.lower()

    if file_ext == '.docx':
        if parsing_method == 'pandoc':
            logger.info(f"尝试使用 Pandoc 方法加载Word文档 {file_path}...")
            return _load_docx_via_pandoc(file_path)
        elif parsing_method == 'direct':
            logger.info(f"使用直接结构化解析方法加载Word文档 {file_path}...")
            return _load_docx_structured(file_path)
        else:
            raise ValueError(f"不支持的解析方法: {parsing_method}")
            
    elif file_ext == '.md':
        logger.info(f"加载Markdown文档 {file_path}...")
        with open(file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()
        return _parse_markdown_text(md_text, file_path.name)
        
    elif file_ext == '.txt':
        logger.info(f"加载纯文本文档 {file_path}...")
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        # 对于纯文本，我们创建一个单一的文档块
        metadata = {"file_name": file_path.name, "block_type": "text"}
        return [Document(text=text, metadata=metadata)]
        
    else:
        raise ValueError(f"不支持的文件类型: {file_ext}")

def _load_docx_structured(file_path: Path) -> List[Document]:
    """
    使用内部逻辑直接结构化解析DOCX文件
    """
    try:
        import docx
        doc = docx.Document(file_path)
        structured_blocks = parse_docx_structure(str(file_path))
        
        if not structured_blocks:
            logger.warning("直接结构化解析未能提取任何块，尝试加载整个文本")
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if full_text:
                return [Document(text=full_text, metadata={"file_name": file_path.name, "block_type": "text"})]
            else:
                return []
                
        documents = [Document(text=block["text"], metadata={
            "file_name": file_path.name,
            "block_type": block.get("type", "unknown"),
            "level": block.get("level", 0),
            "title": block.get("title", ""),
            "is_virtual": block.get("is_virtual", False)
        }) for block in structured_blocks]
        
        logger.info(f"成功使用直接结构化解析加载了 {len(documents)} 个文档节点")
        return documents
    except ImportError:
        logger.error("未安装python-docx库，无法使用直接结构化方法解析Word文档")
        raise ImportError("请安装 python-docx 库: pip install python-docx")
    except Exception as e:
        logger.exception(f"直接解析DOCX文件时出错: {e}")
        # 出错时尝试加载整个文本作为回退
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if full_text:
                logger.warning("直接解析失败，已加载整个文档文本作为回退")
                return [Document(text=full_text, metadata={"file_name": file_path.name, "block_type": "text"})]
        except Exception as fallback_e:
            logger.error(f"回退加载整个文本也失败: {fallback_e}")
        return []

def _load_docx_via_pandoc(file_path: Path) -> List[Document]:
    """
    使用Pandoc将DOCX转换为Markdown，然后解析Markdown
    """
    pandoc_path = shutil.which('pandoc')
    if not pandoc_path:
        logger.warning("系统中未找到 Pandoc。将回退到直接解析方法。")
        logger.warning("请安装 Pandoc (https://pandoc.org/installing.html) 以使用此功能。")
        return _load_docx_structured(file_path)
    
    try:
        logger.info(f"正在使用 Pandoc 将 {file_path} 转换为 Markdown...")
        # 使用 Pandoc 转换为 Markdown CommonMark 格式
        # --wrap=none 避免自动换行破坏段落
        # --markdown-headings=atx 使用 # 风格标题
        process = subprocess.run(
            [pandoc_path, str(file_path), "-f", "docx", "-t", "commonmark", "--wrap=none", "--markdown-headings=atx"],
            capture_output=True,
            text=True,
            encoding='utf-8',
            check=True
        )
        md_text = process.stdout
        logger.info(f"Pandoc 转换完成，Markdown 文本长度: {len(md_text)} 字符")
        
        # 解析生成的 Markdown 文本
        return _parse_markdown_text(md_text, file_path.name)
        
    except subprocess.CalledProcessError as e:
        logger.error(f"Pandoc 执行失败: {e}")
        logger.error(f"Pandoc 返回码: {e.returncode}")
        logger.error(f"Pandoc Stderr: {e.stderr}")
        logger.warning("Pandoc 转换失败，将回退到直接解析方法。")
        return _load_docx_structured(file_path)
    except FileNotFoundError:
        logger.error(f"Pandoc 命令未找到，即使 shutil.which 返回了路径: {pandoc_path}")
        logger.warning("Pandoc 转换失败，将回退到直接解析方法。")
        return _load_docx_structured(file_path)
    except Exception as e:
        logger.exception(f"使用 Pandoc 处理文档时发生意外错误: {e}")
        logger.warning("Pandoc 转换失败，将回退到直接解析方法。")
        return _load_docx_structured(file_path)

def _parse_markdown_text(md_text: str, file_name: str) -> List[Document]:
    """
    解析Markdown文本并提取结构化块
    
    支持多种Markdown标题格式：
    1. 标准格式: # 标题
    2. Pandoc格式: **标题** 或 数字. <span>**标题**
    3. 中文编号格式: 一、二、三、...
    4. 数字层级格式: 1.1、1.2、2.1、...
    5. 括号编号格式: (1)、(2)、（1）、（2）...
    
    新增功能:
    - 提取目录中的链接
    - 建立目录项与标题之间的关联
    """
    import re
    
    if not mistune:
        logger.error("Markdown 解析需要 mistune 库。请运行: pip install mistune")
        # 简单回退：将整个文本视为一个块
        return [Document(text=md_text, metadata={"file_name": file_name, "block_type": "text"})]
    
    try:
        # 检测文档标题 (通常是第一行或带有特殊格式的行)
        document_title = ""
        lines = md_text.strip().split('\n')
        if lines and len(lines) > 0:
            first_line = lines[0].strip()
            if re.match(r'^\*\*.+\*\*$', first_line):  # 匹配 **标题** 格式
                document_title = re.sub(r'^\*\*(.+)\*\*$', r'\1', first_line)
                
        # 提取目录部分
        toc_content = ""
        toc_start_index = -1
        toc_end_index = -1
        toc_links = {}  # 存储目录链接关系: {链接文本: 链接目标}
        
        for i, line in enumerate(lines):
            if line.strip() == "目录":
                toc_start_index = i
                break
        
        if toc_start_index >= 0:
            # 寻找目录结束位置 - 通常是连续的链接行之后的第一个非链接/非空行
            in_links_section = False
            consecutive_non_link_lines = 0
            
            for i in range(toc_start_index + 1, len(lines)):
                line = lines[i].strip()
                
                # 空行不影响判断
                if not line:
                    continue
                    
                # 检查是否是目录项链接并提取链接关系
                is_link_line = bool(re.search(r'\[.*?\]\(.*?\)', line))
                
                if is_link_line:
                    # 提取链接文本和目标
                    link_matches = re.findall(r'\[(.*?)\]\((.*?)\)', line)
                    for link_text, link_target in link_matches:
                        # 存储链接关系，移除链接文本中的前导空格和数字
                        clean_link_text = re.sub(r'^[\s\d\.\(\)（）一二三四五六七八九十]+[、\.\s]*', '', link_text.strip())
                        toc_links[clean_link_text] = link_target
                    
                    in_links_section = True
                    consecutive_non_link_lines = 0
                elif in_links_section:
                    consecutive_non_link_lines += 1
                    
                # 如果连续两个非链接行且已经处于链接区域，则认为目录结束
                if in_links_section and consecutive_non_link_lines >= 2:
                    toc_end_index = i - consecutive_non_link_lines
                    break
                    
                # 如果遇到明确的标题标记，也认为目录结束
                if (re.match(r'^#+\s+', line) or 
                    re.match(r'^<span id=".+?".*?>.*?\*\*.*?\*\*', line) or
                    (re.match(r'^\d+\.\s+', line) and not is_link_line) or
                    re.match(r'^（\d+）\s+', line) or 
                    re.match(r'^\(\d+\)\s+', line) or
                    re.match(r'^[一二三四五六七八九十]+[、\s]', line)):
                    
                    toc_end_index = i
                    break
            
            # 如果没有找到明确的结束位置，设定一个最大范围
            if toc_end_index < 0 and in_links_section:
                # 寻找连续超过3个空行的位置
                empty_line_count = 0
                for i in range(toc_start_index + 1, min(toc_start_index + 200, len(lines))):
                    if not lines[i].strip():
                        empty_line_count += 1
                    else:
                        empty_line_count = 0
                        
                    if empty_line_count >= 3:
                        toc_end_index = i - 2  # 减去多余的空行
                        break
            
            # 如果仍未找到结束位置，设置一个合理的默认值
            if toc_end_index < 0:
                # 设置默认结束位置为起始后的50行或第一个看起来像正文的行
                for i in range(toc_start_index + 1, min(toc_start_index + 50, len(lines))):
                    line = lines[i].strip()
                    # 如果找到一个看起来像段落开始的行（不是链接、不是编号）
                    if (line and not re.search(r'\[.*?\]\(.*?\)', line) and 
                        not re.match(r'^[0-9一二三四五六七八九十]+[、\.]', line) and
                        not re.match(r'^（\d+）', line) and not re.match(r'^\(\d+\)', line)):
                        toc_end_index = i
                        break
                
                # 如果还是没找到，就使用一个固定的结束位置
                if toc_end_index < 0:
                    toc_end_index = min(toc_start_index + 50, len(lines))
            
            if toc_end_index > toc_start_index:
                toc_content = "\n".join(lines[toc_start_index:toc_end_index])
        
        # 处理章节内容
        documents = []
        current_heading = None
        current_level = 0
        current_content = []
        heading_stack = []  # 用于跟踪章节层级
        
        # 如果找到了文档标题，添加为第一个文档节点
        if document_title:
            documents.append(Document(
                text=document_title,
                metadata={
                    "file_name": file_name,
                    "block_type": "document_title",
                    "level": 0,
                    "title": document_title,
                    "breadcrumb_path": document_title
                }
            ))
        
        # 如果找到了目录，添加为文档节点，并包含提取的链接关系
        if toc_content:
            # 将 toc_links 字典序列化为 JSON 字符串以符合 ChromaDB 要求
            toc_links_json = json.dumps(toc_links, ensure_ascii=False) 
            documents.append(Document(
                text=toc_content,
                metadata={
                    "file_name": file_name,
                    "block_type": "toc",
                    "level": 0,
                    "title": "目录",
                    "breadcrumb_path": "目录",
                    "toc_links_json": toc_links_json  # 存储 JSON 字符串
                }
            ))
        
        # 用于存储标题到节点ID的映射
        title_to_id = {}
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 跳过已经处理过的目录部分
            if toc_start_index <= i < toc_end_index:
                i += 1
                continue
            
            # 检测标题 - 支持多种格式
            heading_level = 0
            heading_text = ""
            is_sub_item = False  # 标记是否为子项目（带-前缀）
            
            # 0. 检测带"-"前缀的子标题 (优先检查，这通常是子项)
            if re.match(r'^-\s*(.+)$', line.strip()):
                match = re.match(r'^-\s*(.+)$', line.strip())
                heading_text = match.group(1).strip()
                if heading_text.startswith("**") and heading_text.endswith("**"):
                    heading_text = heading_text[2:-2]
                
                # 暂时标记为2级，后面会根据上下文重新调整
                heading_level = 2
                is_sub_item = True
            
            # 1. 检测标准Markdown标题 (# 标题)
            elif re.match(r'^(#+)\s+(.*)', line):
                heading_match = re.match(r'^(#+)\s+(.*)', line)
                heading_level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
            
            # 2. 检测Pandoc格式标题 - 独立的粗体行 (**标题**)
            elif re.match(r'^\*\*(.+)\*\*$', line.strip()):
                heading_text = re.sub(r'^\*\*(.+)\*\*$', r'\1', line.strip())
                heading_level = 1  # 默认为一级标题
            
            # 3. 检测中文数字编号标题 (一、标题)
            elif re.match(r'^[一二三四五六七八九十]+[、\s]+(.+)$', line.strip()):
                match = re.match(r'^[一二三四五六七八九十]+[、\s]+(.+)$', line.strip())
                heading_text = match.group(1).strip()
                if heading_text.startswith("**") and heading_text.endswith("**"):
                    heading_text = heading_text[2:-2]
                heading_level = 1
            
            # 4. 检测数字层级格式 (1. 标题 或 1.1 标题)
            elif re.match(r'^(\d+)(\.\d+)*\.\s+(.+)$', line.strip()):
                match = re.match(r'^(\d+)(\.\d+)*\.\s+(.+)$', line.strip())
                heading_text = match.group(3).strip()
                if heading_text.startswith("**") and heading_text.endswith("**"):
                    heading_text = heading_text[2:-2]
                
                # 根据点的数量确定级别
                numbering = match.group(1)
                if match.group(2):
                    # 如果有子章节数字 (例如 1.2.3)
                    heading_level = 1 + match.group(2).count('.')
                else:
                    heading_level = 1
            
            # 5. 检测括号编号格式 ((1) 标题 或 （1）标题)
            elif re.match(r'^[（\(]\s*\d+\s*[）\)]\s+(.+)$', line.strip()):
                match = re.match(r'^[（\(]\s*\d+\s*[）\)]\s+(.+)$', line.strip())
                heading_text = match.group(1).strip()
                if heading_text.startswith("**") and heading_text.endswith("**"):
                    heading_text = heading_text[2:-2]
                heading_level = 2  # 通常是二级标题
            
            # 6. 检测Pandoc格式标题 - 带编号和span标签的标题
            elif re.search(r'<span id=".*?".*?>.*?\*\*(.+?)\*\*', line):
                heading_text = re.search(r'\*\*(.+?)\*\*', line).group(1)
                
                # 根据编号形式判断级别
                if re.match(r'^\s*\d+\.\s+', line):  # 例如: "1. <span>**标题**"
                    heading_level = 1
                elif re.match(r'^\s*（\d+）\s+', line) or re.match(r'^\s*\(\d+\)\s+', line):  # 例如: "（1） <span>**标题**"
                    heading_level = 2
                else:
                    heading_level = 1  # 默认为一级标题
            
            # 如果找到了新标题，保存之前的内容
            if heading_level > 0 and heading_text:
                # 保存上一个块
                if current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        # 构建面包屑路径
                        breadcrumb_path = current_heading
                        parent_sections = []
                        if heading_stack:
                            for level, title in heading_stack:
                                if level < current_level:
                                    parent_sections.append(title)
                            
                            if parent_sections:
                                breadcrumb_path = " > ".join(parent_sections + [current_heading])
                        
                        # 将 parent_sections 列表序列化为 JSON 字符串
                        parent_sections_json = json.dumps(parent_sections, ensure_ascii=False)
                        
                        metadata = {
                            "file_name": file_name,
                            "block_type": "section" if current_heading else "text",
                            "level": current_level,
                            "title": current_heading if current_heading else "",
                            "breadcrumb_path": breadcrumb_path,
                            "parent_sections_json": parent_sections_json, # 存储 JSON 字符串
                            "is_sub_item": is_sub_item
                        }
                        
                        # 检查当前标题是否在目录链接中，如果是，添加链接信息
                        if current_heading is not None:
                            clean_title = re.sub(r'^[\s\d\.\(\)（）一二三四五六七八九十]+[、\.\s]*', '', current_heading)
                            if clean_title in toc_links:
                                metadata["toc_link_target"] = toc_links[clean_title]
                        
                        doc = Document(
                            text=f"{current_heading}\n\n{text}" if current_heading else text, 
                            metadata=metadata
                        )
                        documents.append(doc)
                        
                        # 存储标题到文档ID的映射
                        if current_heading:
                            if current_heading is not None:
                                clean_title = re.sub(r'^[\s\d\.\(\)（）一二三四五六七八九十]+[、\.\s]*', '', current_heading)
                                title_to_id[clean_title] = len(documents) - 1  # 使用文档索引作为ID
                
                # 更新heading_stack
                # 移除当前级别或更高级别的标题，除非是子项
                if not is_sub_item:
                    while heading_stack and heading_stack[-1][0] >= heading_level:
                        heading_stack.pop()
                # 添加当前标题到栈中
                heading_stack.append((heading_level, heading_text))
                
                # 开始新块
                current_heading = heading_text
                current_level = heading_level
                current_content = []
            else:
                # 将当前行添加到当前内容块
                current_content.append(line)
            
            i += 1
        
        # 保存最后一个块
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                # 构建面包屑路径
                breadcrumb_path = current_heading
                parent_sections = []
                if heading_stack:
                    for level, title in heading_stack:
                        if level < current_level:
                            parent_sections.append(title)
                    
                    if parent_sections:
                        breadcrumb_path = " > ".join(parent_sections + [current_heading])
                
                # 将 parent_sections 列表序列化为 JSON 字符串
                parent_sections_json = json.dumps(parent_sections, ensure_ascii=False)
                
                metadata = {
                    "file_name": file_name,
                    "block_type": "section" if current_heading else "text",
                    "level": current_level,
                    "title": current_heading if current_heading else "",
                    "breadcrumb_path": breadcrumb_path,
                    "parent_sections_json": parent_sections_json, # 存储 JSON 字符串
                    "is_sub_item": False
                }
                
                # 检查当前标题是否在目录链接中，如果是，添加链接信息
                if current_heading is not None:
                    clean_title = re.sub(r'^[\s\d\.\(\)（）一二三四五六七八九十]+[、\.\s]*', '', current_heading)
                    if clean_title in toc_links:
                        metadata["toc_reference"] = True
                        metadata["toc_link"] = toc_links[clean_title]
                
                doc = Document(
                    text=f"{current_heading}\n\n{text}" if current_heading else text, 
                    metadata=metadata
                )
                documents.append(doc)
                
                # 存储标题到文档ID的映射
                if current_heading:
                    if current_heading is not None:
                        clean_title = re.sub(r'^[\s\d\.\(\)（）一二三四五六七八九十]+[、\.\s]*', '', current_heading)
                        title_to_id[clean_title] = len(documents) - 1  # 使用文档索引作为ID
        
        # 如果没有识别到任何文档块，则返回整个文档作为一个块
        if not documents and md_text:
            documents.append(Document(
                text=md_text,
                metadata={
                    "file_name": file_name,
                    "block_type": "text"
                }
            ))

        # 对文档块进行后处理，确定正确的层级关系
        documents = _post_process_sections(documents)
        
        # 最后处理：建立目录项与节点的双向链接关系
        if len(documents) > 1 and title_to_id:
            # 找到目录节点
            toc_node = None
            for doc in documents:
                if doc.metadata.get("block_type") == "toc":
                    toc_node = doc
                    break
            
            if toc_node:
                # 将 title_to_id 字典序列化为 JSON 字符串
                title_to_id_json = json.dumps(title_to_id, ensure_ascii=False)
                # 在目录节点中添加标题到文档ID的映射 (JSON格式)
                toc_node.metadata["title_to_id_json"] = title_to_id_json
                
                # 对每个节点，添加与目录的关联信息
                for doc in documents:
                    if doc.metadata.get("block_type") == "section" and doc.metadata.get("title"):
                        # 尝试获取和清理标题，处理可能为None的情况
                        current_heading = doc.metadata.get("title")
                        clean_title = ""
                        if current_heading is not None:
                            # 使用与创建 title_to_id 时相同的清理逻辑
                            clean_title = re.sub(r'^<span id=\".*?\"></span>', '', current_heading).strip()
                            clean_title = re.sub(r'\\*\\*(.+?)\\*\\*', r'\\1', clean_title) # 移除粗体标记
                            clean_title = re.sub(r'^[\s\\d\\.\\(\\)\uff08\\uff09\\u4e00-\\u9fa5]+[\\u3001\\.\\s]*', '', clean_title) # 移除编号和前缀

                        if clean_title in toc_links:
                            doc.metadata["toc_reference"] = True
                            doc.metadata["toc_link"] = toc_links[clean_title]

        logger.info(f"Markdown 解析完成，生成了 {len(documents)} 个文档节点")
        return documents

    except Exception as e:
        logger.exception(f"解析Markdown文本时出错: {e}")
        # 出错时返回整个文本作为单一文档
        logger.warning("Markdown解析失败，返回整个文本作为一个块")
        return [Document(text=md_text, metadata={"file_name": file_name, "block_type": "text"})]

def _post_process_sections(documents: List[Document]) -> List[Document]:
    """
    对解析好的文档节点进行后处理，调整节点间的层级关系
    
    主要处理：
    1. 标题以"-"开头的节点，找到其正确的父节点
    2. 根据标题内容判断节点的层级关系
    """
    if len(documents) <= 2:  # 仅有文档标题和目录或更少，无需处理
        return documents
    
    # 提取正常章节（非"-"前缀）作为潜在父节点
    main_sections = []
    hyphen_sections = []
    
    # 第一遍：分类节点
    for i, doc in enumerate(documents):
        if doc.metadata.get("block_type") == "section":
            title = doc.metadata.get("title", "")
            if title.startswith("-"):
                hyphen_sections.append((i, doc))
            else:
                # 忽略文档标题和目录
                if doc.metadata.get("level", 0) > 0:
                    main_sections.append((i, doc))
    
    # 第二遍：确定"-"前缀节点的父节点
    parent_section = None
    
    for idx, doc in hyphen_sections:
        # 寻找最近的前面的主节点作为父节点
        title = doc.metadata.get("title", "")
        parent_found = False
        
        # 向前查找最近的主节点
        for main_idx, main_doc in reversed(main_sections):
            if main_idx < idx:  # 只考虑在当前节点之前的主节点
                parent_section = main_doc
                parent_found = True
                break
        
        if parent_found and parent_section:
            # 更新父章节列表并序列化
            # 先获取原始列表，处理可能不是列表的情况
            parent_sections_str = parent_section.metadata.get("parent_sections_json", "[]")
            try:
                existing_parent_sections = json.loads(parent_sections_str)
                if not isinstance(existing_parent_sections, list):
                    existing_parent_sections = []
            except json.JSONDecodeError:
                existing_parent_sections = []
            
            updated_parent_sections = existing_parent_sections + [parent_section.metadata.get("title", "")]
            updated_parent_sections_json = json.dumps(updated_parent_sections, ensure_ascii=False)
            
            # 更新当前节点的元数据
            parent_title = parent_section.metadata.get("title", "")
            parent_level = parent_section.metadata.get("level", 1)
            
            # 更新节点层级 - 比父节点高一级
            doc.metadata["level"] = parent_level + 1
            
            # 更新面包屑路径和序列化后的父章节
            doc.metadata["parent_sections_json"] = updated_parent_sections_json # 存储 JSON 字符串
            # 面包屑路径仍用原始列表构建
            breadcrumb_parts = [str(p) for p in updated_parent_sections] + [doc.metadata.get("title", "")]
            doc.metadata["breadcrumb_path"] = " > ".join(breadcrumb_parts) 
    
    return documents

def chunk_document(documents: List[Document], chunk_size: int = 512, chunk_overlap: int = 50) -> List[TextNode]:
    """
    将加载的文档分块
    """
    # 实现分块逻辑
    pass

    # 返回分块后的文本节点列表
    return [] 